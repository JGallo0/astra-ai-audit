import io
import json
import unicodedata
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


def _is_valid_xml_char(ch: str) -> bool:
    cp = ord(ch)
    return (
        cp == 0x9
        or cp == 0xA
        or cp == 0xD
        or (0x20 <= cp <= 0xD7FF)
        or (0xE000 <= cp <= 0xFFFD)
        or (0x10000 <= cp <= 0x10FFFF)
    )


def sanitize_xml_text(text: Any) -> str:
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\x00", "")
    text = "".join(ch for ch in text if _is_valid_xml_char(ch))
    return text


def safe_str(value: Any) -> str:
    return sanitize_xml_text("" if value is None else str(value))
    
def normalize_result_language(results: List[Dict[str, Any]], lang: str = "en") -> List[Dict[str, Any]]:
    is_en = lang == "en"

    replacements_pt_to_en = {
        "Não conforme": "Non-compliant",
        "Parcialmente conforme": "Partially compliant",
        "Conforme": "Compliant",
        "Não evidenciado": "Not evidenced",
        "Erro de análise": "Analysis error",
        "alto": "high",
        "medio": "medium",
        "baixo": "low",
        "A evidência disponível do projeto não atende adequadamente ao critério metodológico recuperado.": "The available project evidence does not adequately meet the retrieved methodological criterion.",
        "Há evidência parcial, porém ainda faltam elementos documentais e/ou operacionais para robustez metodológica.": "There is partial evidence, but documentary and/or operational elements are still missing for methodological robustness.",
        "Revisar a aderência metodológica e incluir evidências objetivas que atendam ao requisito.": "Review methodological alignment and include objective evidence that satisfies the requirement.",
        "Complementar a documentação e fortalecer a rastreabilidade/evidência do requisito.": "Complete the documentation and strengthen traceability/evidence for the requirement.",
        "Não identificado.": "Not identified.",
        "Não foram identificados riscos críticos imediatos.": "No immediate critical risks were identified.",
    }

    replacements_en_to_pt = {v: k for k, v in replacements_pt_to_en.items()}

    def normalize_text(text: Any) -> str:
        s = safe_str(text)
        if not s:
            return s

        mapping = replacements_pt_to_en if is_en else replacements_en_to_pt

        # substituição exata primeiro
        if s in mapping:
            return mapping[s]

        # substituição parcial simples
        for src, dst in mapping.items():
            s = s.replace(src, dst)

        return s

    normalized = []

    for item in results:
        row = dict(item)

        for field in [
            "status",
            "risk",
            "gap",
            "recommendation",
            "notes",
            "methodology_basis",
            "project_evidence",
            "title",
            "module",
        ]:
            if field in row:
                row[field] = normalize_text(row[field])

        normalized.append(row)

    return normalized

def flatten_markdown_to_lines(text: str) -> List[str]:
    text = sanitize_xml_text(text)
    return [sanitize_xml_text(line).rstrip() for line in text.replace("\r", "").split("\n")]


def build_audit_dataframe(results: List[Dict[str, Any]]) -> pd.DataFrame:
    df = pd.DataFrame(results)

    preferred_order = [
        "requirement_id",
        "module",
        "title",
        "status",
        "risk",
        "score",
        "confidence",
        "project_evidence",
        "methodology_basis",
        "gap",
        "recommendation",
        "notes",
    ]

    if df.empty:
        return pd.DataFrame(columns=preferred_order)

    existing_cols = [c for c in preferred_order if c in df.columns]
    remaining_cols = [c for c in df.columns if c not in existing_cols]
    return df[existing_cols + remaining_cols]


def convert_df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def convert_json_to_bytes(data: Any) -> bytes:
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def docx_from_text(title: str, text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    h = doc.add_paragraph()
    run = h.add_run(safe_str(title))
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")

    for line in flatten_markdown_to_lines(text):
        stripped = sanitize_xml_text(line).strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(sanitize_xml_text(stripped[2:].strip()))
            r.bold = True
            r.font.size = Pt(15)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(sanitize_xml_text(stripped[3:].strip()))
            r.bold = True
            r.font.size = Pt(13)
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(sanitize_xml_text(stripped[4:].strip()))
            r.bold = True
            r.font.size = Pt(11.5)
        elif stripped.startswith("- "):
            doc.add_paragraph(sanitize_xml_text(stripped[2:].strip()), style="List Bullet")
        else:
            doc.add_paragraph(sanitize_xml_text(stripped))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def pdf_from_text(title: str, text: str) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    left = 50
    right = 50
    top = height - 50
    bottom = 50
    usable_width = width - left - right
    y = top

    def new_page():
        nonlocal y
        c.showPage()
        y = top

    def wrap_line(line: str, font_name: str, font_size: int) -> List[str]:
        words = line.split()
        if not words:
            return [""]
        lines = []
        current = words[0]
        for word in words[1:]:
            test = current + " " + word
            if stringWidth(test, font_name, font_size) <= usable_width:
                current = test
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    c.setTitle(safe_str(title))
    c.setFont("Helvetica-Bold", 16)
    c.drawString(left, y, safe_str(title))
    y -= 28

    for raw_line in flatten_markdown_to_lines(text):
        if y < bottom + 30:
            new_page()

        line = sanitize_xml_text(raw_line).strip()
        if not line:
            y -= 10
            continue

        if line.startswith("# "):
            c.setFont("Helvetica-Bold", 15)
            chunks = wrap_line(line[2:].strip(), "Helvetica-Bold", 15)
            for ch in chunks:
                if y < bottom + 25:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 20
        elif line.startswith("## "):
            c.setFont("Helvetica-Bold", 13)
            chunks = wrap_line(line[3:].strip(), "Helvetica-Bold", 13)
            for ch in chunks:
                if y < bottom + 22:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 18
        elif line.startswith("- "):
            c.setFont("Helvetica", 10)
            bullet_text = "• " + line[2:].strip()
            chunks = wrap_line(bullet_text, "Helvetica", 10)
            for ch in chunks:
                if y < bottom + 18:
                    new_page()
                c.drawString(left + 8, y, sanitize_xml_text(ch))
                y -= 14
        else:
            c.setFont("Helvetica", 10)
            chunks = wrap_line(line, "Helvetica", 10)
            for ch in chunks:
                if y < bottom + 18:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 14

    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def matrix_to_docx_bytes(df: pd.DataFrame, title: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run(safe_str(title))
    r.bold = True
    r.font.size = Pt(15)

    doc.add_paragraph("")

    if df.empty:
        doc.add_paragraph("Nenhum item de matriz de conformidade foi gerado.")
    else:
        cols = list(df.columns)
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = safe_str(col)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                cells[i].text = safe_str(row.get(col, ""))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def matrix_to_pdf_bytes(df: pd.DataFrame, title: str) -> bytes:
    text = [f"# {safe_str(title)}", ""]
    if df.empty:
        text.append("Nenhum item de matriz de conformidade foi gerado.")
    else:
        for i, row in df.iterrows():
            text.append(f"## Item {i + 1}")
            for col in df.columns:
                text.append(f"- {safe_str(col)}: {safe_str(row.get(col, ''))}")
            text.append("")
    return pdf_from_text(title, "\n".join(text))


def build_full_audit_text(summary: Dict[str, Any], results: List[Dict[str, Any]]) -> str:
    lines = []
    lines.append("# Auditoria Resumida Isometric")
    lines.append("")
    lines.append(f"- Total de requisitos avaliados: {summary.get('total_requirements', 0)}")
    lines.append(f"- Score geral: {summary.get('overall_score', 0)}%")
    lines.append(f"- Confiança geral: {summary.get('overall_confidence', 0)}%")
    lines.append("")

    lines.append("## Resumo por status")
    for k, v in (summary.get("status_counts", {}) or {}).items():
        lines.append(f"- {safe_str(k)}: {safe_str(v)}")
    lines.append("")

    lines.append("## Resumo por risco")
    for k, v in (summary.get("risk_counts", {}) or {}).items():
        lines.append(f"- {safe_str(k)}: {safe_str(v)}")
    lines.append("")

    lines.append("## Score por módulo")
    for k, v in (summary.get("module_scores", {}) or {}).items():
        lines.append(f"- {safe_str(k)}: {safe_str(v)}")
    lines.append("")

    lines.append("## Matriz detalhada")
    lines.append("")
    for item in results:
        lines.append(f"### {safe_str(item.get('requirement_id', ''))} — {safe_str(item.get('title', ''))}")
        lines.append(f"- Módulo: {safe_str(item.get('module', ''))}")
        lines.append(f"- Status: {safe_str(item.get('status', ''))}")
        lines.append(f"- Risco: {safe_str(item.get('risk', ''))}")
        lines.append(f"- Score: {safe_str(item.get('score', ''))}")
        lines.append(f"- Confiança: {safe_str(item.get('confidence', ''))}")
        lines.append(f"- Gap: {safe_str(item.get('gap', ''))}")
        lines.append(f"- Recomendação: {safe_str(item.get('recommendation', ''))}")
        lines.append("")
    return "\n".join(lines)


def build_full_eligibility_dossier_text(
    project_name: str,
    summary: Dict[str, Any],
    results: List[Dict[str, Any]],
    trails: List[Dict[str, Any]],
) -> str:
    lines = []
    lines.append("# Dossiê de Elegibilidade Metodológica")
    lines.append("")
    lines.append("## Projeto analisado")
    lines.append(safe_str(project_name))
    lines.append("")
    lines.append("## 1. Síntese executiva")
    lines.append(f"- Total de requisitos avaliados: {summary.get('total_requirements', 0)}")
    lines.append(f"- Score geral: {summary.get('overall_score', 0)}%")
    lines.append(f"- Confiança geral: {summary.get('overall_confidence', 0)}%")
    lines.append("")
    for k, v in (summary.get("status_counts", {}) or {}).items():
        lines.append(f"- Status {safe_str(k)}: {safe_str(v)}")
    for k, v in (summary.get("risk_counts", {}) or {}).items():
        lines.append(f"- Risco {safe_str(k)}: {safe_str(v)}")
    lines.append("")

    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for item in results:
        grouped.setdefault(safe_str(item.get("module", "Sem módulo")), []).append(item)

    lines.append("## 2. Avaliação detalhada")
    lines.append("")
    for module, items in grouped.items():
        lines.append(f"### Módulo: {module}")
        lines.append("")
        for item in items:
            lines.append(f"#### {safe_str(item.get('requirement_id', ''))} — {safe_str(item.get('title', ''))}")
            lines.append(f"- Status: {safe_str(item.get('status', ''))}")
            lines.append(f"- Risco: {safe_str(item.get('risk', ''))}")
            lines.append(f"- Score: {safe_str(item.get('score', ''))}")
            lines.append(f"- Confiança: {safe_str(item.get('confidence', ''))}")
            lines.append("")
            lines.append("**Base metodológica**")
            lines.append(safe_str(item.get("methodology_basis", "")) or "Não identificado.")
            lines.append("")
            lines.append("**Evidência do projeto**")
            lines.append(safe_str(item.get("project_evidence", "")) or "Não identificado.")
            lines.append("")
            lines.append("**Gap**")
            lines.append(safe_str(item.get("gap", "")) or "Não identificado.")
            lines.append("")
            lines.append("**Recomendação**")
            lines.append(safe_str(item.get("recommendation", "")) or "Não identificado.")
            lines.append("")
            if item.get("notes"):
                lines.append("**Notas**")
                lines.append(safe_str(item.get("notes", "")))
                lines.append("")

    lines.append("## 3. Trilha técnica resumida")
    lines.append("")
    for i, trail in enumerate(trails, start=1):
        lines.append(f"### Item {i} — {safe_str(trail.get('module', ''))}")
        lines.append("**Queries do projeto**")
        lines.append(safe_str(trail.get("project_query", "")) or "Não identificado.")
        lines.append("")
        lines.append("**Queries da metodologia**")
        lines.append(safe_str(trail.get("methodology_query", "")) or "Não identificado.")
        lines.append("")
    return "\n".join(lines)


def render_sources_block(title: str, sources: List[Dict[str, Any]], show_attributes: bool, show_snippets: bool):
    st.markdown(f"#### {title}")

    if not sources:
        st.info("Nenhuma fonte foi retornada.")
        return

    for i, src in enumerate(sources, start=1):
        page_display = src.get("page") or "não identificada"
        label = f"{i}. {src['filename']} — pág./seção: {page_display}"

        with st.expander(label):
            st.write(f"**Grupo:** {src.get('source_group')}")
            st.write(f"**Score:** {src.get('score')}")
            st.write(f"**File ID:** {src.get('file_id')}")
            if show_attributes and src.get("attributes"):
                st.json(src["attributes"])
            if show_snippets and src.get("text"):
                st.code(src["text"], language="text")


def badge_html(text: str, kind: str) -> str:
    class_name = {
        "success": "badge-success",
        "warning": "badge-warning",
        "danger": "badge-danger",
        "info": "badge-info",
    }.get(kind, "badge-info")
    return f'<span class="auditoria-badge {class_name}">{safe_str(text)}</span>'


def status_badge(status: str) -> str:
    s = safe_str(status)
    if s == "Conforme":
        return badge_html(s, "success")
    if s in {"Parcialmente conforme", "Não evidenciado"}:
        return badge_html(s, "warning")
    if s in {"Não conforme", "Erro de análise", "Inconsistência documental"}:
        return badge_html(s, "danger")
    return badge_html(s, "info")

def risk_badge(risk: str) -> str:
    r = safe_str(risk).lower()

    if r in ["baixo", "low"]:
        return badge_html("low", "success")

    if r in ["medio", "medium"]:
        return badge_html("medium", "warning")

    if r in ["alto", "high"]:
        return badge_html("high", "danger")

    return badge_html(risk, "info")
    
def build_executive_dossier_text(
    project_name,
    summary,
    results,
    lang="en",
):
    score = summary.get("overall_score", 0)
    confidence = summary.get("overall_confidence", 0)
    risk_counts = summary.get("risk_counts", {}) or {}

    high_risk = risk_counts.get("alto", 0)

    # =========================
    # MATURITY CLASSIFICATION
    # =========================
    if score >= 75:
        maturity = "High readiness for certification"
        interpretation = "The project demonstrates strong technical structure and consistent alignment with methodological requirements."
    elif score >= 55:
        maturity = "Moderate readiness with relevant gaps"
        interpretation = "The project presents a solid technical foundation but contains gaps that may affect certification eligibility."
    elif score >= 40:
        maturity = "Low readiness – significant risks identified"
        interpretation = "The project is partially structured, with critical gaps that must be addressed prior to certification."
    else:
        maturity = "Not eligible in current state"
        interpretation = "The project does not meet minimum certification requirements at this stage."

    # =========================
    # CRITICAL ITEMS
    # =========================
    critical_items = [
        r for r in results if r.get("risk") in ["alto", "high"]
    ]

    critical_items = sorted(
        critical_items,
        key=lambda x: x.get("score", 0)
    )[:5]

    # =========================
    # RECOMMENDATIONS
    # =========================
    recommendations = []
    for r in critical_items:
        rec = safe_str(r.get("recommendation"))
        if rec:
            recommendations.append(rec)

    recommendations = list(dict.fromkeys(recommendations))[:5]

    # =========================
    # BUILD TEXT
    # =========================
    lines = []

    # HEADER
    lines.append("# CO2mply Carbon Compliance Report")
    lines.append("")
    lines.append(f"## Project: {safe_str(project_name)}")
    lines.append("")

    # EXECUTIVE SUMMARY
    lines.append("## Executive Summary")
    lines.append("")
    lines.append(interpretation)
    lines.append("")
    lines.append(f"- Overall Score: {score:.1f}%")
    lines.append(f"- Confidence Level: {confidence:.1f}%")
    lines.append(f"- Certification Readiness: {maturity}")
    lines.append(f"- High Risk Items Identified: {high_risk}")
    lines.append("")

    # RISKS
    lines.append("## Key Risks and Gaps")
    lines.append("")

    if not critical_items:
        lines.append("No critical risks identified.")
    else:
        for r in critical_items:
            lines.append(
                f"- {safe_str(r.get('title'))}: {safe_str(r.get('gap'))}"
            )

    lines.append("")

    # RECOMMENDATIONS
    lines.append("## Priority Actions")
    lines.append("")

    if not recommendations:
        lines.append("No recommendations generated.")
    else:
        for rec in recommendations:
            lines.append(f"- {rec}")

    lines.append("")

    # MODULE SUMMARY
    lines.append("## Module Performance Overview")
    lines.append("")

    modules = {}
    for r in results:
        modules.setdefault(r.get("module", "General"), []).append(r)

    for module, items in modules.items():
        avg_score = sum(i.get("score", 0) for i in items) / len(items)
        lines.append(f"### {module}")
        lines.append(f"- Average Score: {round(avg_score,1)}%")
        lines.append("")

    return "\n".join(lines)
def pdf_from_text_branded(title: str, text: str, brand_name: str = "CO2mply") -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    left = 50
    right = 50
    top = height - 50
    bottom = 50
    usable_width = width - left - right
    y = top

    def new_page():
        nonlocal y
        c.showPage()
        draw_header_footer()
        y = top - 30

    def draw_header_footer():
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left, height - 28, safe_str(brand_name))
        c.setFont("Helvetica", 8)
        c.drawRightString(width - right, height - 28, "Carbon Compliance Intelligence")
        c.line(left, height - 34, width - right, height - 34)

        c.line(left, bottom - 8, width - right, bottom - 8)
        c.setFont("Helvetica", 8)
        c.drawString(left, bottom - 20, safe_str(brand_name))
        c.drawRightString(width - right, bottom - 20, f"Page {c.getPageNumber()}")

    def wrap_line(line: str, font_name: str, font_size: int) -> List[str]:
        words = line.split()
        if not words:
            return [""]
        lines = []
        current = words[0]
        for word in words[1:]:
            test = current + " " + word
            if stringWidth(test, font_name, font_size) <= usable_width:
                current = test
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    c.setTitle(safe_str(title))
    draw_header_footer()

    y = top - 35

    c.setFont("Helvetica-Bold", 20)
    c.drawString(left, y, safe_str(title))
    y -= 28

    c.setFont("Helvetica", 11)
    c.drawString(left, y, f"Generated by {safe_str(brand_name)}")
    y -= 24

    c.setFont("Helvetica", 10)
    c.drawString(left, y, "Confidential report")
    y -= 28

    for raw_line in flatten_markdown_to_lines(text):
        if y < bottom + 30:
            new_page()

        line = sanitize_xml_text(raw_line).strip()
        if not line:
            y -= 8
            continue

        if line.startswith("# "):
            c.setFont("Helvetica-Bold", 15)
            chunks = wrap_line(line[2:].strip(), "Helvetica-Bold", 15)
            for ch in chunks:
                if y < bottom + 25:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 20
        elif line.startswith("## "):
            c.setFont("Helvetica-Bold", 13)
            chunks = wrap_line(line[3:].strip(), "Helvetica-Bold", 13)
            for ch in chunks:
                if y < bottom + 22:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 18
        elif line.startswith("### "):
            c.setFont("Helvetica-Bold", 11)
            chunks = wrap_line(line[4:].strip(), "Helvetica-Bold", 11)
            for ch in chunks:
                if y < bottom + 20:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 16
        elif line.startswith("- "):
            c.setFont("Helvetica", 10)
            bullet_text = "• " + line[2:].strip()
            chunks = wrap_line(bullet_text, "Helvetica", 10)
            for ch in chunks:
                if y < bottom + 18:
                    new_page()
                c.drawString(left + 8, y, sanitize_xml_text(ch))
                y -= 14
        else:
            c.setFont("Helvetica", 10)
            chunks = wrap_line(line, "Helvetica", 10)
            for ch in chunks:
                if y < bottom + 18:
                    new_page()
                c.drawString(left, y, sanitize_xml_text(ch))
                y -= 14

    c.save()
    buffer.seek(0)
    return buffer.getvalue()
